import logging
import re
import tempfile
import os
from io import BytesIO
from docx import Document
import pdfplumber

# Initialize logger
logger = logging.getLogger(__name__)

def extract_text_from_file(file):
    """
    Extract text content from uploaded file based on file type
    
    Args:
        file: Flask uploaded file object
        
    Returns:
        Extracted text content
    """
    filename = file.filename.lower()
    content = file.read()
    
    if filename.endswith('.pdf'):
        return extract_text_from_pdf(BytesIO(content))
    elif filename.endswith('.docx'):
        return extract_text_from_docx(BytesIO(content))
    elif filename.endswith('.txt'):
        return content.decode('utf-8', errors='ignore')
    else:
        logger.warning(f"Unsupported file format: {filename}")
        return ""

def extract_text_from_pdf(file_stream):
    """
    Extract text from PDF file
    
    Args:
        file_stream: File-like object containing PDF data
        
    Returns:
        Extracted text
    """
    try:
        text = ""
        with pdfplumber.open(file_stream) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
        
        # Clean up text
        text = re.sub(r'\s+', ' ', text)
        return text.strip()
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        return ""

def extract_text_from_docx(file_stream):
    """
    Extract text from DOCX file
    
    Args:
        file_stream: File-like object containing DOCX data
        
    Returns:
        Extracted text
    """
    try:
        doc = Document(file_stream)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        return '\n'.join(full_text)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        return ""
